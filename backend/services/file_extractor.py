"""
File content extraction service for BRICK OS.

Supports:
- 3D CAD files (STL, STEP, OBJ, FBX, GLTF, GLB): 100MB max
- PDF documents: 50MB max
- Images (with OCR): 20MB max
- Word/Excel/CSV: 20MB max
- Text/code files: 10MB max
"""

import os
import struct
from typing import Dict, Any, Optional
import logging

logger = logging.getLogger(__name__)


# File size limits by category (bytes)
FILE_SIZE_LIMITS = {
    '3d': 100 * 1024 * 1024,       # 100MB for 3D files
    'pdf': 50 * 1024 * 1024,        # 50MB for PDFs
    'image': 20 * 1024 * 1024,      # 20MB for images
    'document': 20 * 1024 * 1024,   # 20MB for documents
    'text': 10 * 1024 * 1024,       # 10MB for text/code
}

# File type categories
FILE_CATEGORIES = {
    '3d': ['.stl', '.step', '.stp', '.obj', '.fbx', '.gltf', '.glb', '.3mf', '.ply'],
    'pdf': ['.pdf'],
    'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg'],
    'document': ['.docx', '.xlsx', '.xls', '.csv', '.pptx'],
    'text': ['.txt', '.md', '.json', '.yaml', '.yml', '.xml',
             '.py', '.js', '.ts', '.jsx', '.tsx', '.c', '.cpp',
             '.h', '.hpp', '.java', '.go', '.rs', '.swift', '.kt', '.cs']
}


def get_file_category(ext: str) -> str:
    """Get category for file extension."""
    ext = ext.lower()
    for category, extensions in FILE_CATEGORIES.items():
        if ext in extensions:
            return category
    return 'text'  # Default


def get_size_limit(ext: str) -> int:
    """Get size limit for file extension."""
    category = get_file_category(ext)
    return FILE_SIZE_LIMITS.get(category, 10 * 1024 * 1024)


async def extract_file_content(file_path: str, content_type: Optional[str] = None) -> Dict[str, Any]:
    """
    Extract content and metadata from various file types.
    
    Returns:
        Dict with:
        - content: Extracted text/content
        - metadata: File-specific metadata (dims, etc.)
        - category: File category
        - success: Whether extraction succeeded
    """
    ext = os.path.splitext(file_path)[1].lower()
    category = get_file_category(ext)
    
    result = {
        "content": "",
        "metadata": {"category": category, "extension": ext},
        "category": category,
        "success": True,
        "error": None
    }
    
    try:
        if category == '3d':
            # Extract 3D metadata
            result["metadata"].update(await extract_3d_metadata(file_path, ext))
            result["content"] = format_3d_summary(result["metadata"])
            
        elif category == 'image':
            # OCR for images
            result["content"] = await extract_image_ocr(file_path)
            # Get image dimensions
            result["metadata"]["dimensions"] = get_image_dimensions(file_path)
            
        elif category == 'pdf':
            result["content"] = await extract_pdf(file_path)
            
        elif category == 'document':
            if ext == '.docx':
                result["content"] = await extract_docx(file_path)
            elif ext in ['.xlsx', '.xls', '.csv']:
                result["content"] = await extract_spreadsheet(file_path)
            else:
                result["content"] = f"[Document: {os.path.basename(file_path)}]"
                
        else:
            # Text files
            result["content"] = await extract_text(file_path)
            
    except Exception as e:
        logger.error(f"Failed to extract {file_path}: {e}")
        result["success"] = False
        result["error"] = str(e)
        result["content"] = f"[Error extracting file: {str(e)}]"
    
    return result


async def extract_3d_metadata(file_path: str, ext: str) -> Dict[str, Any]:
    """
    Extract metadata from 3D CAD files.
    
    Returns:
        Dict with format, dimensions, volume, triangle_count, etc.
    """
    metadata = {
        "type": "3d_model",
        "format": ext.upper().lstrip('.'),
        "dimensions": None,
        "volume_mm3": None,
        "surface_area_mm2": None,
        "triangle_count": None,
        "vertex_count": None,
        "is_ascii": None
    }
    
    try:
        if ext == '.stl':
            metadata.update(parse_stl(file_path))
        elif ext in ['.step', '.stp']:
            metadata.update(parse_step(file_path))
        elif ext == '.obj':
            metadata.update(parse_obj(file_path))
        elif ext == '.ply':
            metadata.update(parse_ply(file_path))
        else:
            # For other formats, just get file stats
            metadata["file_size"] = os.path.getsize(file_path)
            
    except Exception as e:
        logger.warning(f"Failed to parse 3D file {file_path}: {e}")
        metadata["parse_error"] = str(e)
    
    return metadata


def parse_stl(file_path: str) -> Dict[str, Any]:
    """Parse STL file (binary or ASCII) and extract metadata."""
    metadata = {"format": "STL"}
    
    with open(file_path, 'rb') as f:
        header = f.read(80)
        
        # Check if binary or ASCII
        is_ascii = b'solid' in header.lower()
        metadata["is_ascii"] = is_ascii
        
        if not is_ascii:
            # Binary STL
            f.seek(80)  # Skip header
            num_triangles = struct.unpack('<I', f.read(4))[0]
            metadata["triangle_count"] = num_triangles
            
            # Read first triangle to check validity
            if num_triangles > 0:
                try:
                    # Each triangle: 12 bytes normal + 12 bytes v1 + 12 bytes v2 + 12 bytes v3 + 2 bytes attribute
                    triangle_data = f.read(50)
                    if len(triangle_data) == 50:
                        metadata["valid_binary"] = True
                except:
                    metadata["valid_binary"] = False
        else:
            # ASCII STL - count 'facet normal' occurrences
            f.seek(0)
            content = f.read().decode('utf-8', errors='ignore')
            metadata["triangle_count"] = content.lower().count('facet normal')
            metadata["is_ascii"] = True
    
    # Calculate bounding box from file
    try:
        dims = calculate_stl_dimensions(file_path, is_ascii)
        if dims:
            metadata["dimensions"] = dims
    except Exception as e:
        logger.warning(f"Could not calculate STL dimensions: {e}")
    
    return metadata


def calculate_stl_dimensions(file_path: str, is_ascii: bool) -> Optional[Dict[str, float]]:
    """Calculate bounding box dimensions of STL file."""
    min_coords = [float('inf'), float('inf'), float('inf')]
    max_coords = [float('-inf'), float('-inf'), float('-inf')]
    
    with open(file_path, 'rb') as f:
        if is_ascii:
            f.seek(0)
            content = f.read().decode('utf-8', errors='ignore')
            import re
            # Find all vertex lines: "vertex x y z"
            vertices = re.findall(r'vertex\s+([-\d.]+)\s+([-\d.]+)\s+([-\d.]+)', content)
            for v in vertices:
                x, y, z = float(v[0]), float(v[1]), float(v[2])
                min_coords[0] = min(min_coords[0], x)
                min_coords[1] = min(min_coords[1], y)
                min_coords[2] = min(min_coords[2], z)
                max_coords[0] = max(max_coords[0], x)
                max_coords[1] = max(max_coords[1], y)
                max_coords[2] = max(max_coords[2], z)
        else:
            # Binary STL
            f.seek(80 + 4)  # Skip header and triangle count
            import numpy as np
            
            while True:
                try:
                    # Read triangle (normal + 3 vertices)
                    data = f.read(50)  # 12*4 + 2 = 50 bytes per triangle
                    if len(data) < 50:
                        break
                    
                    # Skip normal (12 bytes), read 3 vertices
                    for i in range(3):
                        offset = 12 + (i * 12)
                        x, y, z = struct.unpack('<fff', data[offset:offset+12])
                        min_coords[0] = min(min_coords[0], x)
                        min_coords[1] = min(min_coords[1], y)
                        min_coords[2] = min(min_coords[2], z)
                        max_coords[0] = max(max_coords[0], x)
                        max_coords[1] = max(max_coords[1], y)
                        max_coords[2] = max(max_coords[2], z)
                        
                except Exception:
                    break
    
    if min_coords[0] == float('inf'):
        return None
    
    return {
        "x_mm": round(max_coords[0] - min_coords[0], 3),
        "y_mm": round(max_coords[1] - min_coords[1], 3),
        "z_mm": round(max_coords[2] - min_coords[2], 3),
        "x_in": round((max_coords[0] - min_coords[0]) / 25.4, 3),
        "y_in": round((max_coords[1] - min_coords[1]) / 25.4, 3),
        "z_in": round((max_coords[2] - min_coords[2]) / 25.4, 3),
        "bounding_box": {
            "min": [round(c, 3) for c in min_coords],
            "max": [round(c, 3) for c in max_coords]
        }
    }


def parse_step(file_path: str) -> Dict[str, Any]:
    """Parse STEP file and extract metadata."""
    metadata = {"format": "STEP"}
    
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
        
        # Count entities
        metadata["entity_count"] = content.count('#')
        
        # Look for product name
        import re
        product_match = re.search(r"#\d+\s*=\s*PRODUCT\s*\([^,]*,\s*'([^']+)'", content)
        if product_match:
            metadata["product_name"] = product_match.group(1)
        
        # Look for shape representation
        if 'MANIFOLD_SOLID_BREP' in content or 'CLOSED_SHELL' in content:
            metadata["has_solid_geometry"] = True
        
        # Try to find dimensions in content
        dim_matches = re.findall(r"(\d+\.?\d*)\s*([Mm][Mm]|[Ii][Nn])", content)
        if dim_matches:
            metadata["mentioned_dimensions"] = dim_matches[:5]  # First 5
    
    return metadata


def parse_obj(file_path: str) -> Dict[str, Any]:
    """Parse Wavefront OBJ file and extract metadata."""
    metadata = {"format": "OBJ"}
    
    vertex_count = 0
    normal_count = 0
    texcoord_count = 0
    face_count = 0
    
    min_coords = [float('inf'), float('inf'), float('inf')]
    max_coords = [float('-inf'), float('-inf'), float('-inf')]
    
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        for line in f:
            line = line.strip()
            if line.startswith('v '):
                vertex_count += 1
                parts = line.split()
                if len(parts) >= 4:
                    x, y, z = float(parts[1]), float(parts[2]), float(parts[3])
                    min_coords[0] = min(min_coords[0], x)
                    min_coords[1] = min(min_coords[1], y)
                    min_coords[2] = min(min_coords[2], z)
                    max_coords[0] = max(max_coords[0], x)
                    max_coords[1] = max(max_coords[1], y)
                    max_coords[2] = max(max_coords[2], z)
            elif line.startswith('vn '):
                normal_count += 1
            elif line.startswith('vt '):
                texcoord_count += 1
            elif line.startswith('f '):
                face_count += 1
    
    metadata["vertex_count"] = vertex_count
    metadata["normal_count"] = normal_count
    metadata["texcoord_count"] = texcoord_count
    metadata["face_count"] = face_count
    
    if min_coords[0] != float('inf'):
        metadata["dimensions"] = {
            "x_mm": round(max_coords[0] - min_coords[0], 3),
            "y_mm": round(max_coords[1] - min_coords[1], 3),
            "z_mm": round(max_coords[2] - min_coords[2], 3)
        }
    
    return metadata


def parse_ply(file_path: str) -> Dict[str, Any]:
    """Parse PLY file and extract metadata."""
    metadata = {"format": "PLY"}
    
    with open(file_path, 'rb') as f:
        header_lines = []
        while True:
            line = f.readline().decode('utf-8', errors='ignore').strip()
            header_lines.append(line)
            if line == 'end_header':
                break
        
        header = '\n'.join(header_lines)
        
        # Parse header
        import re
        vertex_match = re.search(r'element vertex (\d+)', header)
        face_match = re.search(r'element face (\d+)', header)
        
        if vertex_match:
            metadata["vertex_count"] = int(vertex_match.group(1))
        if face_match:
            metadata["face_count"] = int(face_match.group(1))
        
        metadata["is_ascii"] = 'format ascii' in header.lower()
    
    return metadata


def format_3d_summary(metadata: Dict[str, Any]) -> str:
    """Format 3D metadata as readable text summary."""
    lines = [f"[3D Model: {metadata.get('format', 'Unknown')}]"]
    
    dims = metadata.get('dimensions')
    if dims:
        lines.append(f"Dimensions: {dims.get('x_mm', 0)} x {dims.get('y_mm', 0)} x {dims.get('z_mm', 0)} mm")
        lines.append(f"({dims.get('x_in', 0)} x {dims.get('y_in', 0)} x {dims.get('z_in', 0)} inches)")
    
    if metadata.get('triangle_count'):
        lines.append(f"Triangles: {metadata['triangle_count']:,}")
    if metadata.get('vertex_count'):
        lines.append(f"Vertices: {metadata['vertex_count']:,}")
    if metadata.get('face_count'):
        lines.append(f"Faces: {metadata['face_count']:,}")
    
    return '\n'.join(lines)


async def extract_image_ocr(file_path: str) -> str:
    """Extract text from image using OCR."""
    try:
        from PIL import Image
        
        image = Image.open(file_path)
        
        # Get basic info
        width, height = image.size
        format_name = image.format or "Unknown"
        mode = image.mode
        
        description = f"[Image: {width}x{height} pixels, {format_name}, {mode}]"
        
        # Try OCR
        try:
            import pytesseract
            text = pytesseract.image_to_string(image)
            if text.strip():
                return f"{description}\n\nExtracted text:\n{text}"
            return description
        except ImportError:
            return f"{description}\n[OCR not available]"
            
    except Exception as e:
        return f"[Image error: {str(e)}]"


def get_image_dimensions(file_path: str) -> Optional[Dict[str, int]]:
    """Get image dimensions."""
    try:
        from PIL import Image
        with Image.open(file_path) as img:
            return {"width": img.width, "height": img.height}
    except:
        return None


async def extract_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber or PyPDF2."""
    # Try pdfplumber first (better extraction)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "pages": len(pdf.pages),
                "title": pdf.metadata.get('Title', ''),
                "author": pdf.metadata.get('Author', '')
            }
            
            for i, page in enumerate(pdf.pages[:20]):  # Limit to 20 pages
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        
        text = '\n\n'.join(text_parts)
        if text.strip():
            header = f"[PDF: {metadata['pages']} pages]"
            if metadata['title']:
                header += f"\nTitle: {metadata['title']}"
            return f"{header}\n\n{text}"
        
        return "[PDF: No text content found]"
        
    except ImportError:
        pass  # Fall through to PyPDF2
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            
            for i, page in enumerate(reader.pages[:20]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i+1} ---\n{page_text}")
                except:
                    continue
        
        text = '\n\n'.join(text_parts)
        return f"[PDF: {num_pages} pages]\n\n{text}" if text.strip() else f"[PDF: {num_pages} pages, no text]"
        
    except ImportError:
        return "[PDF: Install pdfplumber or PyPDF2 to extract text]"


async def extract_docx(file_path: str) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs[:100]:  # Limit to first 100 paragraphs
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables (convert to text)
        for table in doc.tables[:5]:  # Limit to first 5 tables
            text_parts.append("\n[Table]")
            for row in table.rows[:10]:  # Limit rows per table
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))
        
        text = '\n'.join(text_parts)
        return f"[DOCX]\n\n{text}" if text.strip() else "[DOCX: Empty or no text content]"
        
    except ImportError:
        return "[DOCX: python-docx not installed]"


async def extract_spreadsheet(file_path: str) -> str:
    """Extract text from Excel/CSV as markdown table."""
    try:
        import pandas as pd
        
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path, nrows=100)  # Limit to 100 rows
        else:
            df = pd.read_excel(file_path, nrows=100)
        
        # Convert to markdown table
        markdown = df.to_markdown(index=False)
        return f"[Spreadsheet: {len(df)} rows x {len(df.columns)} columns]\n\n{markdown}"
        
    except ImportError:
        return "[Spreadsheet: pandas not installed]"
    except Exception as e:
        return f"[Spreadsheet error: {str(e)}]"


async def extract_text(file_path: str) -> str:
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            
        ext = os.path.splitext(file_path)[1]
        lines = content.split('\n')
        
        # Add file info header
        header = f"[Text file: {len(lines)} lines, {len(content)} chars]"
        
        # For code files, include more context
        if ext in ['.py', '.js', '.ts', '.c', '.cpp', '.java', '.go']:
            # Count functions/classes
            import re
            functions = len(re.findall(r'\bdef\s+\w+|\bfunction\s+\w+|\bclass\s+\w+', content))
            header += f", ~{functions} functions/classes"
        
        return f"{header}\n\n{content}"
        
    except Exception as e:
        return f"[Text error: {str(e)}]"
